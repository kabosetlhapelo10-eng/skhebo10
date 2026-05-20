from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
MARKDOWN = ROOT / "WEBSITE_SPECIFICATION_REPORT.md"
OUTPUT = ROOT / "Motion_Clothing_Website_Report.docx"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text.strip())
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def is_table_line(line):
    return line.startswith("|") and line.endswith("|")


def parse_table(lines, start):
    rows = []
    index = start
    while index < len(lines) and is_table_line(lines[index].strip()):
        row = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in row):
            rows.append(row)
        index += 1
    return rows, index


def add_table(document, rows):
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r, row in enumerate(rows):
        for c in range(len(table.columns)):
            cell = table.cell(r, c)
            text = row[c] if c < len(row) else ""
            set_cell_text(cell, text, bold=(r == 0))
            if r == 0:
                shade_cell(cell, "111111")
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
    document.add_paragraph()


def add_wireframe_images(document):
    screenshots = [
        ("Desktop wireframe/page template", "screenshots/Desktop - 1 (6).png"),
        ("Android expanded wireframe/page template", "screenshots/Android Expanded - 1 (3).png"),
        ("Android medium wireframe/page template", "screenshots/Android Medium - 1 (10).png"),
    ]
    for caption, shot in screenshots:
        path = ROOT / shot
        if not path.exists():
            continue
        paragraph = document.add_paragraph(caption)
        paragraph.runs[0].bold = True
        document.add_picture(str(path), width=Inches(5.8))


def add_code_block(document, code_lines):
    paragraph = document.add_paragraph()
    paragraph.style = "No Spacing"
    for line in code_lines:
        run = paragraph.add_run(line + "\n")
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def add_markdown(document, text):
    lines = text.splitlines()
    index = 0
    in_code = False
    code_lines = []

    while index < len(lines):
        raw = lines[index]
        line = raw.strip()

        if line.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue

        if in_code:
            code_lines.append(raw)
            index += 1
            continue

        if not line:
            index += 1
            continue

        if is_table_line(line):
            rows, index = parse_table(lines, index)
            add_table(document, rows)
            continue

        if line.startswith("# "):
            paragraph = document.add_heading(line[2:], level=0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            index += 1
            continue

        if line.startswith("## "):
            document.add_heading(line[3:], level=1)
            index += 1
            continue

        if line.startswith("### "):
            document.add_heading(line[4:], level=2)
            if line == "### Wireframe Evidence":
                index += 1
                while index < len(lines) and lines[index].strip():
                    paragraph = document.add_paragraph(lines[index].strip().lstrip("- "), style="List Bullet" if lines[index].strip().startswith("- ") else None)
                    index += 1
                add_wireframe_images(document)
                continue
            index += 1
            continue

        if line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
            index += 1
            continue

        if re.match(r"^\d+\.\s", line):
            document.add_paragraph(re.sub(r"^\d+\.\s", "", line), style="List Number")
            index += 1
            continue

        paragraph = document.add_paragraph()
        for part in re.split(r"(\*\*.*?\*\*)", line):
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)
        index += 1


def style_document(document):
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Title"]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.color.rgb = RGBColor(17, 17, 17)
    for section in document.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)


def add_screenshots(document):
    screenshots = [
        "screenshots/Desktop - 1 (6).png",
        "screenshots/Android Expanded - 1 (3).png",
        "screenshots/Android Medium - 1 (10).png",
    ]
    document.add_heading("Appendix: QA Screenshot Evidence", level=1)
    for shot in screenshots:
        path = ROOT / shot
        if not path.exists():
            continue
        document.add_paragraph(path.name)
        document.add_picture(str(path), width=Inches(5.8))


def main():
    document = Document()
    style_document(document)
    add_markdown(document, MARKDOWN.read_text(encoding="utf-8"))
    add_screenshots(document)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
